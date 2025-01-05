import tkinter as tk
from tkinter import messagebox
from docx import Document
import datetime
from tabulate import tabulate


class Kendaraan:
    def __init__(self, nomor_kendaraan, jenis_kendaraan, nama_kendaraan, nama_pemilik):
        self.nomor_kendaraan = nomor_kendaraan
        self.jenis_kendaraan = jenis_kendaraan
        self.nama_kendaraan = nama_kendaraan
        self.nama_pemilik = nama_pemilik
        self.tanggal_masuk = datetime.datetime.now().strftime('%d-%m-%Y')
        self.waktu_masuk = datetime.datetime.now().strftime('%H:%M:%S')
        self.tanggal_keluar = None
        self.waktu_keluar = None

    def hitung_biaya(self, jam):
        pass

class Sepeda(Kendaraan):
    def hitung_biaya(self, jam):
        return 2000 * jam

class Motor(Kendaraan):
    def hitung_biaya(self, jam):
        return 4000 * jam

class Mobil(Kendaraan):
    def hitung_biaya(self, jam):
        return 6000 * jam

class Sistem_Parkir:
    def __init__(self):
        self.kendaraan = []


    def masuk_kendaraan(self, kendaraan):
        kendaraan.tanggal_masuk = datetime.datetime.now().strftime('%d-%m-%Y')
        kendaraan.waktu_masuk = datetime.datetime.now().strftime('%H:%M:%S')
        self.kendaraan.append(kendaraan)
        messagebox.showinfo("Info", "Detail rekaman tersimpan")

        receipt = CetakReceipt()
        receipt.cetak_receipt(kendaraan)

    def hapus_entri(self, nomor_kendaraan):
        for kendaraan in self.kendaraan:
            if kendaraan.nomor_kendaraan == nomor_kendaraan:
                self.kendaraan.remove(kendaraan)
                messagebox.showinfo("Info", "Dihapus dengan sukses")
                break
        else:
            messagebox.showinfo("Info", "Tidak Ada Entri Seperti Itu")

    def lihat_kendara_terparkir(self):
        count=0
        data = []
        for kendaraan in self.kendaraan:
            count += 1
            data.append([
                kendaraan.nomor_kendaraan,
                kendaraan.jenis_kendaraan,
                kendaraan.nama_kendaraan,
                kendaraan.nama_pemilik,
                kendaraan.tanggal_masuk,
                kendaraan.waktu_masuk
            ])

        headers = ["Nomor Kendaraan", "Jenis Kendaraan", "Nama Kendaraan", "Nama Pemilik", "Tanggal", "Waktu"]

        result = tabulate(data, headers, tablefmt="pretty")
        result += f"\nTotal Rekaman = {count}"

        messagebox.showinfo("Kendaraan Terparkir", result)

    def lihat_sisa_tempat_parkir(self, sepeda, motor, mobil):
        result = ""
        result += "--------------------------------------------------------------------------------------------------------------------------------------------------------------\n"
        result += "\t\tTempat Parkir Tersisa\n"
        result += "--------------------------------------------------------------------------------------------------------------------------------------------------------------\n"
        sisa_sepeda = sum(1 for v in self.kendaraan if isinstance(v, Sepeda))
        sisa_motor = sum(1 for v in self.kendaraan if isinstance(v, Motor))
        sisa_mobil = sum(1 for v in self.kendaraan if isinstance(v, Mobil))
        result += f"\tTempat Tersedia untuk Sepeda - {sepeda - sisa_sepeda}\n"
        result += f"\tTempat Tersedia untuk Motor - {motor - sisa_motor}\n"
        result += f"\tTempat Tersedia untuk Mobil - {mobil - sisa_mobil}\n"
        result += "--------------------------------------------------------------------------------------------------------------------------------------------------------------\n"
        messagebox.showinfo("Sisa Tempat Parkir", result)

    def buat_tagihan(self, nomor_kendaraan):
        for kendaraan in self.kendaraan:
            if kendaraan.nomor_kendaraan == nomor_kendaraan:
                if kendaraan.tanggal_keluar is None or kendaraan.waktu_keluar is None:
                    kendaraan.tanggal_keluar = datetime.datetime.now().strftime('%d-%m-%Y')
                    kendaraan.waktu_keluar = datetime.datetime.now().strftime('%H:%M:%S')

                jam_masuk = datetime.datetime.strptime(f"{kendaraan.tanggal_masuk} {kendaraan.waktu_masuk}", "%d-%m-%Y %H:%M:%S")
                jam_keluar = datetime.datetime.strptime(f"{kendaraan.tanggal_keluar} {kendaraan.waktu_keluar}", "%d-%m-%Y %H:%M:%S")
                selisih_jam = (jam_keluar - jam_masuk).seconds // 3600

                biaya = kendaraan.hitung_biaya(selisih_jam)
                biaya_tambahan = 0.18 * biaya
                total_biaya = biaya + biaya_tambahan
                result = ""
                result += f"\tWaktu Check-in Kendaraan - {kendaraan.waktu_masuk}\n"
                result += f"\tTanggal Check-in Kendaraan - {kendaraan.tanggal_masuk}\n"
                result += f"\tWaktu Check-out Kendaraan - {kendaraan.waktu_keluar}\n"
                result += f"\tTanggal Check-out Kendaraan - {kendaraan.tanggal_keluar}\n"
                result += f"\tJenis Kendaraan - {kendaraan.jenis_kendaraan}\n"
                result += f"\tBiaya Parkir - {biaya}\n"
                result += f"\tBiaya Tambahan 18% - {biaya_tambahan}\n"
                result += f"\tTotal Biaya - {total_biaya}\n"
                result += "\n-----Terima kasih telah menggunakan layanan kami-----\n"
                messagebox.showinfo("Tagihan Parkir", result)
                break
        else:
            messagebox.showinfo("Info", "Tidak Ada Entri Seperti Itu")


class CetakReceipt(Sistem_Parkir):
    @staticmethod
    def cetak_receipt(kendaraan):
        document = Document()
        document.add_heading('Karcis', 0)

        document.add_paragraph(f'Nomor Kendaraan: {kendaraan.nomor_kendaraan}')
        document.add_paragraph(f'Jenis Kendaraan: {kendaraan.jenis_kendaraan}')
        document.add_paragraph(f'Nama Kendaraan: {kendaraan.nama_kendaraan}')
        document.add_paragraph(f'Nama Pemilik: {kendaraan.nama_pemilik}')
        document.add_paragraph(f'Tanggal: {kendaraan.tanggal_masuk}')
        document.add_paragraph(f'Waktu: {kendaraan.waktu_masuk}')

        document.save('karcis.docx')
        messagebox.showinfo("Karcis", "Karcis berhasil dicetak (karcis.docx)")

class AplikasiParkir(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Aplikasi Manajemen Parkir")
        self.geometry("600x400")

        self.sistem_parkir = Sistem_Parkir()

        self.label = tk.Label(self, text="Manajemen Sistem Parkir", font=("Helvetica", 16))
        self.label.pack(pady=10)

        self.button_masuk_kendaraan = tk.Button(self, text="Masuk Kendaraan", command=self.masuk_kendaraan)
        self.button_masuk_kendaraan.pack(pady=5)

        self.button_hapus_entri = tk.Button(self, text="Hapus Entri", command=self.hapus_entri)
        self.button_hapus_entri.pack(pady=5)

        self.button_lihat_kendaraan_terparkir = tk.Button(self, text="Lihat Kendaraan Terparkir", command=self.lihat_kendara_terparkir)
        self.button_lihat_kendaraan_terparkir.pack(pady=5)

        self.button_lihat_sisa_tempat_parkir = tk.Button(self, text="Lihat Sisa Tempat Parkir", command=self.lihat_sisa_tempat_parkir)
        self.button_lihat_sisa_tempat_parkir.pack(pady=5)

        self.button_buat_tagihan = tk.Button(self, text="Buat Tagihan", command=self.buat_tagihan)
        self.button_buat_tagihan.pack(pady=5)

        self.button_keluar = tk.Button(self, text="Keluar", command=self.destroy)
        self.button_keluar.pack(pady=5)

    def masuk_kendaraan(self):
        dialog_masuk_kendaraan = DialogMasukKendaraan(self)
        self.wait_window(dialog_masuk_kendaraan)
        kendaraan = dialog_masuk_kendaraan.result
        if kendaraan:
            self.sistem_parkir.masuk_kendaraan(kendaraan)
            self.set_waktu_masuk()  # Panggil metode set_waktu_masuk

    def set_waktu_masuk(self):
        self.waktu_masuk = datetime.datetime.now().strftime('%H:%M:%S')

    def hapus_entri(self):
        dialog_hapus_entri = DialogHapusEntri(self)
        self.wait_window(dialog_hapus_entri)
        nomor_kendaraan = dialog_hapus_entri.result
        if nomor_kendaraan:
            self.sistem_parkir.hapus_entri(nomor_kendaraan)

    def lihat_kendara_terparkir(self):
        self.sistem_parkir.lihat_kendara_terparkir()

    def lihat_sisa_tempat_parkir(self):
        self.sistem_parkir.lihat_sisa_tempat_parkir(78, 100, 250)

    def buat_tagihan(self):
        dialog_buat_tagihan = DialogBuatTagihan(self)
        self.wait_window(dialog_buat_tagihan)
        nomor_kendaraan, selisih_jam = dialog_buat_tagihan.result
        if nomor_kendaraan:
            self.sistem_parkir.buat_tagihan(nomor_kendaraan)

    def get_jam_parkir(self):
        if hasattr(self, 'waktu_masuk') and self.waktu_masuk:
            waktu_keluar = datetime.datetime.now().strftime('%H:%M:%S')
            jam_masuk = datetime.datetime.strptime(self.waktu_masuk, "%H:%M:%S")
            jam_keluar = datetime.datetime.strptime(waktu_keluar, "%H:%M:%S")
            selisih_jam = (jam_keluar - jam_masuk).seconds // 3600
            return selisih_jam
        else:
            messagebox.showwarning("Peringatan", "Waktu masuk kendaraan belum diatur.")
            return 0

class DialogMasukKendaraan(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)

        self.title("Masuk Kendaraan")
        self.geometry("300x200")

        self.result = None

        self.label_nomor_kendaraan = tk.Label(self, text="Nomor Kendaraan:")
        self.label_nomor_kendaraan.pack(pady=5)

        self.entry_nomor_kendaraan = tk.Entry(self)
        self.entry_nomor_kendaraan.pack(pady=5)

        self.label_jenis_kendaraan = tk.Label(self, text="Jenis Kendaraan:")
        self.label_jenis_kendaraan.pack(pady=5)

        self.entry_jenis_kendaraan = tk.Entry(self)
        self.entry_jenis_kendaraan.pack(pady=5)

        self.label_nama_kendaraan = tk.Label(self, text="Nama Kendaraan:")
        self.label_nama_kendaraan.pack(pady=5)

        self.entry_nama_kendaraan = tk.Entry(self)
        self.entry_nama_kendaraan.pack(pady=5)

        self.label_nama_pemilik = tk.Label(self, text="Nama Pemilik:")
        self.label_nama_pemilik.pack(pady=5)

        self.entry_nama_pemilik = tk.Entry(self)
        self.entry_nama_pemilik.pack(pady=5)

        self.button_ok = tk.Button(self, text="OK", command=self.ok)
        self.button_ok.pack(pady=10)

    def ok(self):
        nomor_kendaraan = self.entry_nomor_kendaraan.get().upper()
        jenis_kendaraan = self.entry_jenis_kendaraan.get().lower()
        nama_kendaraan = self.entry_nama_kendaraan.get()
        nama_pemilik = self.entry_nama_pemilik.get()

        if jenis_kendaraan not in ('a', 'b', 'c'):
            messagebox.showwarning("Peringatan", "Jenis kendaraan tidak valid.")
            return

        if jenis_kendaraan == 'a':
            kendaraan = Sepeda(nomor_kendaraan, "Sepeda", nama_kendaraan, nama_pemilik)
        elif jenis_kendaraan == 'b':
            kendaraan = Motor(nomor_kendaraan, "Motor", nama_kendaraan, nama_pemilik)
        elif jenis_kendaraan == 'c':
            kendaraan = Mobil(nomor_kendaraan, "Mobil", nama_kendaraan, nama_pemilik)

        self.result = kendaraan
        self.destroy()

class DialogHapusEntri(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)

        self.title("Hapus Entri")
        self.geometry("250x150")

        self.result = None

        self.label_nomor_kendaraan = tk.Label(self, text="Nomor Kendaraan:")
        self.label_nomor_kendaraan.pack(pady=5)

        self.entry_nomor_kendaraan = tk.Entry(self)
        self.entry_nomor_kendaraan.pack(pady=5)

        self.button_ok = tk.Button(self, text="OK", command=self.ok)
        self.button_ok.pack(pady=10)

    def ok(self):
        nomor_kendaraan = self.entry_nomor_kendaraan.get().upper()
        self.result = nomor_kendaraan
        self.destroy()

class DialogBuatTagihan(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)

        self.title("Buat Tagihan")
        self.geometry("300x150")

        self.result = None

        self.label_nomor_kendaraan = tk.Label(self, text="Nomor Kendaraan:")
        self.label_nomor_kendaraan.pack(pady=5)

        self.entry_nomor_kendaraan = tk.Entry(self)
        self.entry_nomor_kendaraan.pack(pady=5)

        self.button_ok = tk.Button(self, text="OK", command=self.ok)
        self.button_ok.pack(pady=10)

    def ok(self):
        nomor_kendaraan = self.entry_nomor_kendaraan.get().upper()
        waktu_masuk = self.master.waktu_masuk  # Ambil waktu masuk dari sistem parkir

        # Ubah waktu_masuk menjadi objek datetime
        waktu_masuk = datetime.datetime.strptime(waktu_masuk, '%H:%M:%S')

        # Hitung selisih jam
        waktu_sekarang = datetime.datetime.now()
        selisih_jam = int((waktu_sekarang - waktu_masuk).seconds // 3600)

        self.result = nomor_kendaraan, selisih_jam
        self.destroy()

if __name__ == "__main__":
    app = AplikasiParkir()
    app.mainloop()